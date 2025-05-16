# maintenance/maintenance_planner.py
import pandas as pd
import os
from datetime import datetime, timedelta
import random
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import _Cell
import re

class MaintenancePlanner:
    """
    Classe pour générer des gammes de maintenance à partir des données AMDEC
    """
    
    def __init__(self, amdec_df=None):
        """
        Initialisation avec les données AMDEC si disponibles
        
        Args:
            amdec_df (pandas.DataFrame, optional): DataFrame contenant les données AMDEC.
                Si non fourni, une valeur par défaut sera utilisée.
        """
        self.amdec_df = amdec_df
        self.maintenance_data = {}
        self.output_path = None
        
        # Mappings des composants et sous-composants pour les chemins d'images
        self.image_mappings = {
            'economiseur bt': {
                'collecteur sortie': ['eco_bt_collecteur_1.png', 'eco_bt_collecteur_2.png'],
                'epingle': ['eco_bt_epingle_1.png', 'eco_bt_epingle_2.png']
            },
            'economiseur ht': {
                'collecteur entree': ['eco_ht_collecteur_1.png', 'eco_ht_collecteur_2.png'],
                'tubes suspension': ['eco_ht_tubes_1.png', 'eco_ht_tubes_2.png']
            },
            'surchauffeur bt': {
                'epingle': ['sur_bt_epingle_1.png', 'sur_bt_epingle_2.png'],
                'collecteur entree': ['sur_bt_collecteur_1.png', 'sur_bt_collecteur_2.png']
            },
            'surchauffeur ht': {
                'tube porteur': ['sur_ht_tube_1.png', 'sur_ht_tube_2.png'],
                'branches entree': ['sur_ht_branches_1.png', 'sur_ht_branches_2.png'],
                'collecteur sortie': ['sur_ht_collecteur_1.png', 'sur_ht_collecteur_2.png']
            },
            'rechauffeur bt': {
                'collecteur entree': ['rch_bt_collecteur_1.png', 'rch_bt_collecteur_2.png'],
                'tubes suspension': ['rch_bt_tubes_1.png', 'rch_bt_tubes_2.png'],
                'tube porteur': ['rch_bt_porteur_1.png', 'rch_bt_porteur_2.png']
            },
            'rechauffeur ht': {
                'branches sortie': ['rch_ht_branches_1.png', 'rch_ht_branches_2.png'],
                'collecteur entree': ['rch_ht_collecteur_1.png', 'rch_ht_collecteur_2.png'],
                'collecteur sortie': ['rch_ht_collecteur_sortie_1.png', 'rch_ht_collecteur_sortie_2.png']
            }
        }
    
    def get_criticality(self, component, subcomponent):
        """
        Récupère la criticité pour un composant et sous-composant spécifiques
        
        Args:
            component (str): Nom du composant
            subcomponent (str): Nom du sous-composant
            
        Returns:
            int: Criticité ou valeur par défaut si non trouvée
        """
        if self.amdec_df is None:
            # Si pas de données AMDEC, retourner une valeur par défaut
            return self._get_default_criticality(component, subcomponent)
        
        # Normaliser les noms pour la recherche
        component = component.lower()
        subcomponent = subcomponent.lower()
        
        # Convertir les colonnes en minuscules pour la recherche
        amdec_df_lower = self.amdec_df.copy()
        amdec_df_lower['Composant'] = amdec_df_lower['Composant'].str.lower()
        amdec_df_lower['Sous-composant'] = amdec_df_lower['Sous-composant'].str.lower()
        
        # Rechercher des correspondances exactes
        filtered_df = amdec_df_lower[
            (amdec_df_lower['Composant'] == component) & 
            (amdec_df_lower['Sous-composant'] == subcomponent)
        ]
        
        if not filtered_df.empty:
            return filtered_df['C'].max()
        
        # Si pas de correspondance exacte, rechercher des correspondances partielles
        for comp_col in amdec_df_lower['Composant'].unique():
            if component in comp_col or comp_col in component:
                for subcomp_col in amdec_df_lower[amdec_df_lower['Composant'] == comp_col]['Sous-composant'].unique():
                    if subcomponent in subcomp_col or subcomp_col in subcomponent:
                        filtered_df = amdec_df_lower[
                            (amdec_df_lower['Composant'] == comp_col) & 
                            (amdec_df_lower['Sous-composant'] == subcomp_col)
                        ]
                        if not filtered_df.empty:
                            return filtered_df['C'].max()
        
        # Si toujours pas de correspondance, retourner une valeur par défaut
        return self._get_default_criticality(component, subcomponent)
    
    def _get_default_criticality(self, component, subcomponent):
        """
        Retourne une criticité par défaut basée sur le composant et sous-composant
        
        Args:
            component (str): Nom du composant
            subcomponent (str): Nom du sous-composant
            
        Returns:
            int: Criticité par défaut
        """
        # Valeurs par défaut basées sur l'expérience
        default_values = {
            'economiseur bt': {
                'collecteur sortie': 45,
                'epingle': 24
            },
            'economiseur ht': {
                'collecteur entree': 24,
                'tubes suspension': 16
            },
            'surchauffeur bt': {
                'epingle': 40,
                'collecteur entree': 24
            },
            'surchauffeur ht': {
                'tube porteur': 30,
                'branches entree': 24,
                'collecteur sortie': 30
            },
            'rechauffeur bt': {
                'collecteur entree': 30,
                'tubes suspension': 24,
                'tube porteur': 24
            },
            'rechauffeur ht': {
                'branches sortie': 36,
                'collecteur entree': 24,
                'collecteur sortie': 20
            }
        }
        
        component = component.lower()
        subcomponent = subcomponent.lower()
        
        if component in default_values and subcomponent in default_values[component]:
            return default_values[component][subcomponent]
        
        # Valeur par défaut générique
        return 25
    
    def generate_maintenance_plan(self, component, subcomponent, criticality=None):
        """
        Génère un plan de maintenance pour un composant et sous-composant spécifiques
        
        Args:
            component (str): Nom du composant
            subcomponent (str): Nom du sous-composant
            criticality (int, optional): Criticité. Si non fournie, sera calculée.
            
        Returns:
            dict: Données du plan de maintenance
        """
        # Normaliser les noms
        component_norm = component.lower()
        subcomponent_norm = subcomponent.lower()
        
        # Déterminer la criticité si non fournie
        if criticality is None:
            criticality = self.get_criticality(component_norm, subcomponent_norm)
        
        # Mapper les noms normalisés aux noms formatés
        component_formatted = self._format_component_name(component_norm)
        subcomponent_formatted = self._format_subcomponent_name(subcomponent_norm)
        
        # Générer les éléments du plan de maintenance
        material_list = self._generate_material_list(component_norm, subcomponent_norm, criticality)
        operations = self._generate_operations(component_norm, subcomponent_norm, criticality)
        total_time = self._calculate_total_time(operations)
        image_paths = self._get_image_paths(component_norm, subcomponent_norm)
        
        # Créer un dictionnaire avec les données du plan
        self.maintenance_data = {
            'component': component_formatted,
            'subcomponent': subcomponent_formatted,
            'criticality': criticality,
            'material_list': material_list,
            'operations': operations,
            'total_time': total_time,
            'image_paths': image_paths,
            'date': datetime.now().strftime("%d/%m/%Y")
        }
        
        return self.maintenance_data
    
    def _format_component_name(self, component):
        """
        Formate le nom du composant pour l'affichage
        
        Args:
            component (str): Nom du composant à formater
            
        Returns:
            str: Nom formaté
        """
        # Mappings des noms de composants
        component_mappings = {
            'economiseur bt': 'Économiseur BT',
            'economiseur ht': 'Économiseur HT',
            'surchauffeur bt': 'Surchauffeur BT',
            'surchauffeur ht': 'Surchauffeur HT',
            'rechauffeur bt': 'Réchauffeur BT',
            'rechauffeur ht': 'Réchauffeur HT'
        }
        
        return component_mappings.get(component, component.title())
    
    def _format_subcomponent_name(self, subcomponent):
        """
        Formate le nom du sous-composant pour l'affichage
        
        Args:
            subcomponent (str): Nom du sous-composant à formater
            
        Returns:
            str: Nom formaté
        """
  # Mappings des noms de sous-composants
        subcomponent_mappings = {
            'epingle': 'Épingle',
            'collecteur entree': 'Collecteur entrée',
            'collecteur sortie': 'Collecteur sortie',
            'branches entree': 'Branches entrée',
            'branches sortie': 'Branches sortie',
            'tube porteur': 'Tube porteur',
            'tubes suspension': 'Tubes suspension'
        }
        
        return subcomponent_mappings.get(subcomponent, subcomponent.title())
    
    def _generate_material_list(self, component, subcomponent, criticality):
        """
        Génère la liste des matériels nécessaires pour la maintenance
        
        Args:
            component (str): Nom du composant
            subcomponent (str): Nom du sous-composant
            criticality (int): Criticité
            
        Returns:
            list: Liste des matériels
        """
        # Liste de base des matériels communs
        common_materials = [
            'Lampe torche',
            'Caméra d\'inspection',
            'Appareil photo'
        ]
        
        # Matériels spécifiques par type de sous-composant
        specific_materials = {
            'epingle': [
                'Appareil à ultrasons',
                'Gel de contact',
                'Brosse semi-rigide',
                'Produit nettoyant',
                'Poste à souder',
                'Électrodes de soudure'
            ],
            'collecteur': [
                'Kit de test d\'étanchéité',
                'Endoscope',
                'Brosse métallique',
                'Pinceau / rouleau',
                'Peinture anticorrosion',
                'Produit dégraissant'
            ],
            'tube': [
                'Appareil à ultrasons',
                'Gel de contact',
                'Brosse flexible',
                'Capteurs de température',
                'Miroir d\'inspection',
                'Peinture anticorrosion'
            ],
            'branches': [
                'Kit de test d\'étanchéité',
                'Outils de serrage',
                'Brosse métallique',
                'Produit nettoyant',
                'Caméra thermique',
                'Joints de rechange'
            ]
        }
        
        # Déterminer la catégorie de sous-composant
        subcomponent_category = None
        for category in specific_materials.keys():
            if category in subcomponent:
                subcomponent_category = category
                break
        
        # Si aucune catégorie n'est trouvée, utiliser une liste générique
        if subcomponent_category is None:
            materials = common_materials + [
                'Brosse métallique',
                'Produit nettoyant',
                'Caméra endoscopique',
                'Peinture anticorrosion'
            ]
        else:
            # Nombre de matériels spécifiques à ajouter en fonction de la criticité
            if criticality <= 12:
                num_materials = 3
            elif criticality <= 20:
                num_materials = 4
            else:
                num_materials = 5
            
            # Sélectionner les matériels spécifiques
            selected_specific = random.sample(specific_materials[subcomponent_category], 
                                            min(num_materials, len(specific_materials[subcomponent_category])))
            
            # Combiner les listes
            materials = common_materials + selected_specific
        
        return materials
    
    def _generate_operations(self, component, subcomponent, criticality):
        """
        Génère les opérations de maintenance
        
        Args:
            component (str): Nom du composant
            subcomponent (str): Nom du sous-composant
            criticality (int): Criticité
            
        Returns:
            list: Liste des opérations avec leurs détails
        """
        # Opérations de base
        base_operations = [
            {
                'name': 'Inspection visuelle',
                'details': 'Vérification de l\'état général, recherche de signes de corrosion, fuites ou fissures',
                'time': 15  # minutes
            }
        ]
        
        # Opérations spécifiques par type de sous-composant et criticité
        specific_operations = {
            # Pour les épingles
            'epingle': {
                'low': [  # Criticité <= 12
                    {
                        'name': 'Contrôle visuel des soudures',
                        'details': 'Vérification de l\'intégrité des soudures et des raccords',
                        'time': 20
                    }
                ],
                'medium': [  # 12 < Criticité <= 20
                    {
                        'name': 'Contrôle ultrasons',
                        'details': 'Mesure de l\'épaisseur des parois par ultrasons',
                        'time': 25
                    },
                    {
                        'name': 'Nettoyage des dépôts internes',
                        'details': 'Élimination des dépôts et résidus avec une brosse souple',
                        'time': 30
                    }
                ],
                'high': [  # Criticité > 20
                    {
                        'name': 'Réparation ponctuelle',
                        'details': 'Soudure des zones présentant des signes d\'usure ou de faiblesse',
                        'time': 40
                    },
                    {
                        'name': 'Traitement anticorrosion',
                        'details': 'Application d\'un revêtement protecteur sur les surfaces exposées',
                        'time': 35
                    }
                ]
            },
            # Pour les collecteurs
            'collecteur': {
                'low': [
                    {
                        'name': 'Vérification étanchéité',
                        'details': 'Test de pression pour détecter les fuites potentielles',
                        'time': 20
                    }
                ],
                'medium': [
                    {
                        'name': 'Nettoyage interne',
                        'details': 'Élimination des dépôts et de la corrosion interne',
                        'time': 25
                    },
                    {
                        'name': 'Traitement anticorrosion',
                        'details': 'Application d\'un revêtement protecteur',
                        'time': 20
                    }
                ],
                'high': [
                    {
                        'name': 'Inspection endoscopique',
                        'details': 'Examen détaillé des surfaces internes avec caméra',
                        'time': 30
                    },
                    {
                        'name': 'Test ressuage',
                        'details': 'Détection des microfissures par ressuage',
                        'time': 35
                    },
                    {
                        'name': 'Renforcement des zones critiques',
                        'details': 'Traitement spécial des zones soumises à fortes contraintes',
                        'time': 30
                    }
                ]
            },
            # Pour les tubes
            'tube': {
                'low': [
                    {
                        'name': 'Contrôle visuel des points d\'ancrage',
                        'details': 'Vérification des fixations et supports',
                        'time': 15
                    }
                ],
                'medium': [
                    {
                        'name': 'Mesure d\'épaisseur par ultrasons',
                        'details': 'Contrôle de l\'épaisseur des parois aux points critiques',
                        'time': 30
                    },
                    {
                        'name': 'Nettoyage des surfaces externes',
                        'details': 'Élimination des dépôts et résidus de combustion',
                        'time': 25
                    }
                ],
                'high': [
                    {
                        'name': 'Installation capteurs de température',
                        'details': 'Mise en place de capteurs aux points critiques',
                        'time': 35
                    },
                    {
                        'name': 'Analyse vibratoire',
                        'details': 'Mesure des vibrations en fonctionnement',
                        'time': 30
                    },
                    {
                        'name': 'Renforcement structurel',
                        'details': 'Installation de supports supplémentaires',
                        'time': 40
                    }
                ]
            },
            # Pour les branches
            'branches': {
                'low': [
                    {
                        'name': 'Contrôle des raccords',
                        'details': 'Vérification de l\'intégrité des raccordements',
                        'time': 20
                    }
                ],
                'medium': [
                    {
                        'name': 'Test d\'étanchéité',
                        'details': 'Recherche de fuites aux jonctions',
                        'time': 25
                    },
                    {
                        'name': 'Nettoyage interne',
                        'details': 'Élimination des dépôts accumulés',
                        'time': 30
                    }
                ],
                'high': [
                    {
                        'name': 'Remplacement des joints',
                        'details': 'Installation de nouveaux joints d\'étanchéité',
                        'time': 30
                    },
                    {
                        'name': 'Inspection thermographique',
                        'details': 'Détection des points chauds en fonctionnement',
                        'time': 25
                    },
                    {
                        'name': 'Renforcement des raccords',
                        'details': 'Amélioration des points de jonction critiques',
                        'time': 35
                    }
                ]
            }
        }
        
        # Déterminer la catégorie de sous-composant
        subcomponent_category = None
        for category in specific_operations.keys():
            if category in subcomponent:
                subcomponent_category = category
                break
        
        # Si aucune catégorie n'est trouvée, utiliser des opérations génériques
        if subcomponent_category is None:
            operations = base_operations + [
                {
                    'name': 'Contrôle général',
                    'details': 'Vérification complète de l\'état et du fonctionnement',
                    'time': 30
                },
                {
                    'name': 'Nettoyage préventif',
                    'details': 'Élimination des dépôts et contaminants',
                    'time': 25
                }
            ]
        else:
            # Déterminer le niveau de criticité
            if criticality <= 12:
                level = 'low'
            elif criticality <= 20:
                level = 'medium'
            else:
                level = 'high'
            
            # Combiner les opérations
            operations = base_operations.copy()
            
            # Ajouter les opérations de niveau inférieur
            if level in ['medium', 'high']:
                operations.extend(specific_operations[subcomponent_category]['low'])
            
            # Ajouter les opérations du niveau actuel
            operations.extend(specific_operations[subcomponent_category][level])
        
        # Ajouter des numéros d'ordre
        for i, op in enumerate(operations, 1):
            op['order'] = i
        
        return operations
    
    def _calculate_total_time(self, operations):
        """
        Calcule le temps total estimé pour toutes les opérations
        
        Args:
            operations (list): Liste des opérations
            
        Returns:
            str: Temps total formaté en heures et minutes
        """
        total_minutes = sum(op['time'] for op in operations)
        
        hours = total_minutes // 60
        minutes = total_minutes % 60
        
        if hours > 0:
            return f"{hours}h{minutes:02d}min"
        else:
            return f"{minutes}min"
    
    def _get_image_paths(self, component, subcomponent):
        """
        Récupère les chemins des images pour le composant et sous-composant
        
        Args:
            component (str): Nom du composant
            subcomponent (str): Nom du sous-composant
            
        Returns:
            list: Liste des chemins d'images
        """
        # Vérifier si le composant et le sous-composant sont dans les mappings
        if component in self.image_mappings and subcomponent in self.image_mappings[component]:
            image_names = self.image_mappings[component][subcomponent]
            
            # Créer les chemins complets
            image_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'image')
            
            # Vérifier si le dossier existe
            if not os.path.exists(image_dir):
                os.makedirs(image_dir, exist_ok=True)
            
            # Créer des chemins relatifs
            image_paths = [os.path.join('image', img) for img in image_names]
            
            return image_paths
        
        # Retourner une liste vide si pas d'images disponibles
        return []
    
    def save_to_file(self, output_path=None):
        """
        Sauvegarde le plan de maintenance dans un fichier Word
        
        Args:
            output_path (str, optional): Chemin de sortie pour le fichier Word.
                Si non fourni, un chemin par défaut sera utilisé.
        
        Returns:
            str: Chemin du fichier sauvegardé
        """
        if not self.maintenance_data:
            raise ValueError("Aucun plan de maintenance à sauvegarder. Appelez d'abord la méthode generate_maintenance_plan().")
        
        if output_path is None:
            # Créer un répertoire si nécessaire
            maintenance_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'maintenance')
            os.makedirs(maintenance_dir, exist_ok=True)
            
            # Créer un nom de fichier
            component_name = re.sub(r'[^\w\s]', '', self.maintenance_data['component']).replace(' ', '_').lower()
            subcomponent_name = re.sub(r'[^\w\s]', '', self.maintenance_data['subcomponent']).replace(' ', '_').lower()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            output_path = os.path.join(maintenance_dir, f"gamme_{component_name}_{subcomponent_name}.docx")
        
        # Créer un nouveau document Word
        doc = Document()
        
        # Paramètres du document
        sections = doc.sections
        for section in sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
        
        # Titre
        title = f"Gamme de maintenance {self.maintenance_data['component']} / {self.maintenance_data['subcomponent']}"
        doc.add_heading(title, level=1)
        
        # Informations générales
        doc.add_heading(f"{self.maintenance_data['component']} - {self.maintenance_data['subcomponent']} C = {self.maintenance_data['criticality']}", level=2)
        
        # Matériels nécessaires
        doc.add_heading("Matériels nécessaires :", level=3)
        materials_list = doc.add_paragraph()
        materials_list.style = 'List Bullet'
        for material in self.maintenance_data['material_list']:
            materials_list.add_run(f"{material}\n")
        
        # Temps total estimé
        doc.add_heading("Temps total estimé :", level=3)
        doc.add_paragraph(self.maintenance_data['total_time'])
        
        # Tableau des opérations
        doc.add_heading("Opérations à réaliser :", level=2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        
        # En-têtes du tableau
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Ordre"
        hdr_cells[1].text = "Opérations"
        hdr_cells[2].text = "Temps alloué"
        hdr_cells[3].text = "Observations"
        
        # Mettre en forme les en-têtes
        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(11)
        
        # Ajouter les opérations
        for op in self.maintenance_data['operations']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(op['order'])
            row_cells[1].text = f"{op['name']}\n{op['details']}"
            row_cells[2].text = f"{op['time']} min"
            
            # Case vide pour les observations
            row_cells[3].text = ""
            
            # Mettre en forme le texte
            for i, cell in enumerate(row_cells):
                if i == 0:  # Colonne Ordre
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        # Ajouter une mention de sécurité
        doc.add_heading("Consignes de sécurité :", level=3)
        safety_para = doc.add_paragraph()
        safety_para.add_run("• Porter les équipements de protection individuelle (EPI) : casque, lunettes, gants, chaussures de sécurité.\n")
        safety_para.add_run("• Procéder à la consignation complète (électrique, mécanique, thermique) avant toute intervention.\n")
        safety_para.add_run("• Vérifier l'absence de pression et la température avant démontage.\n")
        safety_para.add_run("• Baliser la zone d'intervention.\n")
        
        # Ajouter la date
        doc.add_paragraph(f"Date : {self.maintenance_data['date']}")
        
        # Sauvegarder le document
        doc.save(output_path)
        
        self.output_path = output_path
        return output_path
    
    def get_material_list(self):
        """
        Récupère la liste des matériels nécessaires
        
        Returns:
            list: Liste des matériels
        """
        if not self.maintenance_data:
            raise ValueError("Aucun plan de maintenance généré. Appelez d'abord la méthode generate_maintenance_plan().")
        
        return self.maintenance_data['material_list']
    
    def get_operations(self):
        """
        Récupère la liste des opérations
        
        Returns:
            list: Liste des opérations
        """
        if not self.maintenance_data:
            raise ValueError("Aucun plan de maintenance généré. Appelez d'abord la méthode generate_maintenance_plan().")
        
        return self.maintenance_data['operations']